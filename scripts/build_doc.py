#!/usr/bin/env python3
"""Generate the 'Maintenance Wizard – Understanding & Plan' Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Theme colors ----------
NAVY = RGBColor(0x0B, 0x2B, 0x4A)
STEEL = RGBColor(0x1F, 0x5C, 0x8B)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY = "F2F4F7"
BAND = "0B2B4A"

doc = Document()

# ---------- Base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10.5, white=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "1F5C8B")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = STEEL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(title, text, fill="EAF2F8", barcolor="1F5C8B"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    # left bar
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), barcolor)
    borders.append(left)
    tcPr.append(borders)
    doc.add_paragraph()
    return tbl


def make_table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        shade_cell(hdr[i], BAND)
        set_cell_text(hdr[i], htext, bold=True, white=True, size=10.5)
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            set_cell_text(cells[ci], val, size=10)
            if ri % 2 == 1:
                shade_cell(cells[ci], LIGHT_GREY)
    if widths:
        for row in tbl.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    return tbl


# =========================================================
# COVER
# =========================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TATA STEEL AI HACKATHON 2026")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Round 2 — Agentic AI Challenge")
r.font.size = Pt(12)
r.font.color.rgb = GREY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MAINTENANCE WIZARD")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("for Industrial Equipment")
r.font.size = Pt(18)
r.font.color.rgb = STEEL

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Problem Understanding, Solution Approach,\nArchitecture & Technology Stack")
r.font.size = Pt(13)
r.font.color.rgb = GREY

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("An AI decision-support system that helps maintenance engineers\ndiagnose, predict, prioritise and act — with minimal human effort.")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = STEEL

doc.add_page_break()

# =========================================================
# 0. WHAT THIS DOCUMENT IS
# =========================================================
h1("0. About This Document")
body("This document is the very first step before we start building. Its only job is to make sure "
     "we all understand the problem in plain, simple language before a single line of product code is written.")
body("It answers six questions:")
bullet("What is this problem, really?", bold_lead="1. ")
bullet("What exactly do we have to build?", bold_lead="2. ")
bullet("How will we build it (our approach)?", bold_lead="3. ")
bullet("What will the architecture look like?", bold_lead="4. ")
bullet("What technology stack will we use, and why?", bold_lead="5. ")
bullet("What is our step-by-step plan and what do we have to submit?", bold_lead="6. ")
body("Wherever a technical word appears, it is explained the first time. A full glossary is at the end.")

# =========================================================
# 1. THE PROBLEM IN EASY LANGUAGE
# =========================================================
h1("1. The Problem — In Very Easy Language")

h2("1.1 The everyday situation")
body("A steel plant is full of huge, expensive machines — furnaces, rolling mills, motors, pumps, "
     "conveyors, cranes. These machines run almost non-stop. When one of them suddenly stops working "
     "(this is called \"downtime\"), the whole production line can halt. That costs a lot of money every "
     "single hour, and can even be a safety risk.")

h2("1.2 What a maintenance engineer does today")
body("When a machine misbehaves, the maintenance engineer has to play detective. To figure out what's "
     "wrong and how to fix it, they dig through many separate, scattered sources of information:")
bullet("Thick equipment manuals (PDFs).")
bullet("SOPs — Standard Operating Procedures (the official \"how to do this safely\" steps).")
bullet("Old maintenance logs — records of past repairs.")
bullet("Failure reports — write-ups of what went wrong before.")
bullet("Sensor alerts — warnings from instruments measuring temperature, vibration, pressure, etc.")
body("They read, compare, remember past cases, and use experience to decide what to do. This is slow, "
     "depends heavily on one expert's memory, and two engineers might reach different conclusions for the "
     "same problem.")

callout("In one line:",
        "Maintenance knowledge is scattered everywhere and locked in experts' heads. Finding the right "
        "answer fast is hard. We are building an AI \"wizard\" that brings it all together and gives clear, "
        "trustworthy answers.")

h2("1.3 What the Maintenance Wizard fixes")
body("Our system becomes the engineer's smart assistant. The engineer can simply ask, in normal English "
     "(or type a fault code, or paste an alert), and the Wizard will:")
bullet("Tell them the most likely fault (diagnosis).")
bullet("Explain the probable root cause (why it happened).")
bullet("Predict how much useful life the equipment has left, and warn early if a big failure is coming.")
bullet("Rate how risky and urgent the situation is.")
bullet("Give clear, step-by-step repair actions and a maintenance plan.")
bullet("Show its reasoning and the exact manual/log it used, so the engineer can trust it.")
bullet("Get smarter over time by learning from the engineer's feedback.")

h2("1.4 Why this is an \"Agentic AI\" problem")
body("\"Agentic AI\" means the system does not just answer one question and stop. It can plan a sequence of "
     "steps, use different tools, make decisions, and carry out a multi-step task on its own — like a "
     "junior engineer who knows when to read the manual, when to check the sensor history, and when to "
     "calculate risk. Our solution is built as a team of cooperating AI \"agents\", each an expert at one "
     "job, coordinated by a manager agent. That is exactly what this round is testing.")

# =========================================================
# 2. WHAT WE HAVE TO BUILD
# =========================================================
h1("2. What We Have To Build")

h2("2.1 The inputs (what goes IN)")
make_table(
    ["Category", "Examples (plain language)"],
    [
        ["Operational & failure data", "Equipment delay logs, fault/error codes, failure analysis reports, past breakdown summaries."],
        ["Condition monitoring data", "Sensor summaries (temperature, vibration, pressure), anomaly alerts, process health indicators."],
        ["Knowledge & documents", "Equipment manuals, maintenance SOPs, historical maintenance records, spare-parts info (stock + how long it takes to procure)."],
        ["User interaction", "Plain-English questions, troubleshooting scenarios, and follow-up questions in a back-and-forth chat."],
    ],
    widths=[2.0, 4.6],
)

h2("2.2 The outputs (what comes OUT)")
make_table(
    ["Output group", "What the engineer receives"],
    [
        ["Diagnostic & predictive", "Probable fault, root-cause analysis, Remaining Useful Life (RUL), early warning of catastrophic failure, process-defect detection."],
        ["Risk & priority", "Risk level (Low / Medium / High / Critical), urgency, plant-level bottleneck ranking, priority based on criticality, delay severity, spares availability & procurement lead time."],
        ["Maintenance recommendations", "Step-by-step repair guidance, immediate action points, an optimised maintenance plan, long-term monitoring advice, spare-procurement strategy."],
        ["Reports", "Structured maintenance reports, abnormal-alert reports, decision summaries for supervisors, and an automatic digital logbook entry."],
    ],
    widths=[2.0, 4.6],
)

h2("2.3 The seven things the system MUST do (functional requirements)")
numbered("Use an LLM/SLM to understand questions and reason about maintenance.", bold_lead="Contextual reasoning — ")
numbered("Read and combine manuals, SOPs, history, and logs to answer.", bold_lead="Knowledge integration — ")
numbered("Accept plain-English questions and handle multi-turn conversations.", bold_lead="Natural language chat — ")
numbered("Every answer is traceable to the source document, rule, or record it came from.", bold_lead="Explainable answers — ")
numbered("Detect abnormal behaviour, raise early warnings, and predict failures.", bold_lead="Abnormality detection & prediction — ")
numbered("Let engineers confirm/correct answers so the system improves.", bold_lead="Feedback loop — ")
numbered("Generate live abnormal-alert reports and notify the right people.", bold_lead="Real-time alerting — ")

callout("Note (our LLM choice):",
        "LLM = Large Language Model (the AI that understands and generates text). Our reasoning engine is "
        "OpenAI's GPT, used via API (which the brief explicitly allows). We use a single LLM throughout — "
        "no small/local model — for both agent reasoning and document embeddings.")

h2("2.4 Optional enhancements — WE ARE KEEPING ALL OF THEM")
body("As requested, every optional enhancement is in scope and reflected in the architecture:")
bullet("Conversational chat interface for engineers.")
bullet("Visualisation dashboard for equipment health, trends, and anomalies.")
bullet("Integration with a simulated IoT / equipment-monitoring data stream.")
bullet("A dynamic knowledge base that grows per equipment / system.")
bullet("An automatic digital logbook of maintenance activities and observations.")
bullet("User-role-based alerts and recommendations (engineer vs supervisor vs planner).")

# =========================================================
# 3. HOW WE WILL BUILD IT (APPROACH)
# =========================================================
h1("3. How We Will Build It — Our Approach")

body("The big idea: instead of one giant AI trying to do everything, we build a team of specialised "
     "AI agents coordinated by a manager (\"orchestrator\") agent. This mirrors how a real maintenance "
     "team works and makes each part testable, explainable, and easy to improve.")

h2("3.1 Two backbone techniques")
bullet("the AI doesn't \"guess\" from memory. Before answering, it "
       "retrieves the relevant pages from manuals/SOPs/history and answers using that text — so every "
       "answer is grounded and quotable. This is how we get explainability and trust.",
       bold_lead="RAG (Retrieval-Augmented Generation) — ")
bullet("classic machine-learning + statistics watch the sensor "
       "data to flag abnormal patterns, estimate Remaining Useful Life, and predict failures. The LLM then "
       "explains those numbers in plain language.",
       bold_lead="Predictive analytics & anomaly detection — ")

h2("3.2 The team of agents")
make_table(
    ["Agent", "Its single job"],
    [
        ["Orchestrator (Manager)", "Understands the engineer's request, plans which agents to call and in what order, and assembles the final answer."],
        ["Diagnostic Agent", "Identifies the most probable fault from symptoms, fault codes, and manuals."],
        ["Root-Cause Agent", "Works backward through logs + history to find WHY the fault happened."],
        ["Prediction / RUL Agent", "Uses sensor trends + ML models to estimate remaining life and early failure warnings."],
        ["Anomaly Detection Agent", "Continuously scans incoming sensor data for abnormal behaviour and raises alerts."],
        ["Risk & Priority Agent", "Scores risk level and urgency using criticality, delay, spares stock & lead time."],
        ["Recommendation Agent", "Produces step-by-step actions, the maintenance plan, and spare-procurement strategy."],
        ["Report / Logbook Agent", "Generates structured reports and writes the automatic digital logbook entry."],
        ["Feedback / Learning Agent", "Captures engineer corrections and feeds them back to improve future answers."],
    ],
    widths=[1.9, 4.7],
)

h2("3.3 How a real request flows (a simple story)")
body("An engineer types: \"Mill motor M-204 is tripping repeatedly and vibration alert is high — what do I do?\"")
numbered("The Orchestrator reads the request and plans the steps.")
numbered("The Anomaly + Prediction agents check M-204's recent sensor data and trends.")
numbered("The Diagnostic agent retrieves the motor manual + matching past cases (RAG) and names the likely fault.")
numbered("The Root-Cause agent explains why (e.g., bearing wear → imbalance → vibration → trips).")
numbered("The Risk & Priority agent rates it (e.g., HIGH, because M-204 is a bottleneck and the spare bearing has a 6-day lead time).")
numbered("The Recommendation agent gives immediate steps + a plan + \"order the bearing now\".")
numbered("The Report agent produces a clean summary and a logbook entry; everything cites its sources.")
numbered("The engineer marks it helpful or corrects it — the Feedback agent stores that for next time.")

# =========================================================
# 4. SYSTEM ARCHITECTURE
# =========================================================
h1("4. System Architecture")

h2("4.1 The five layers (text diagram)")
arch = (
    "┌───────────────────────────────────────────────────────────────┐\n"
    "│  1. PRESENTATION LAYER                                          │\n"
    "│     Chat UI  •  Health Dashboard  •  Alerts & Reports  •  Roles │\n"
    "└───────────────────────────────────────────────────────────────┘\n"
    "                              ▲  ▼   (REST / WebSocket API)\n"
    "┌───────────────────────────────────────────────────────────────┐\n"
    "│  2. AGENTIC ORCHESTRATION LAYER  (the \"brain\")                  │\n"
    "│     Orchestrator → Diagnostic • Root-Cause • Prediction/RUL •   │\n"
    "│     Anomaly • Risk/Priority • Recommendation • Report • Feedback│\n"
    "└───────────────────────────────────────────────────────────────┘\n"
    "        ▲                    ▲                      ▲\n"
    "┌──────────────┐   ┌──────────────────┐   ┌──────────────────────┐\n"
    "│ 3. KNOWLEDGE │   │ 4. PREDICTIVE /  │   │   LLM REASONING       │\n"
    "│   (RAG)      │   │   ML ENGINE      │   │  (OpenAI GPT)         │\n"
    "│ Vector store │   │ Anomaly, RUL,    │   │                       │\n"
    "│ manuals/SOPs │   │ failure models   │   │                       │\n"
    "└──────────────┘   └──────────────────┘   └──────────────────────┘\n"
    "                              ▲\n"
    "┌───────────────────────────────────────────────────────────────┐\n"
    "│  5. DATA LAYER                                                  │\n"
    "│  Sensor stream (simulated IoT) • Logs • History DB • Spares •   │\n"
    "│  Feedback store • Digital Logbook                              │\n"
    "└───────────────────────────────────────────────────────────────┘"
)
p = doc.add_paragraph()
run = p.add_run(arch)
run.font.name = "Consolas"
run.font.size = Pt(8)
run.font.color.rgb = NAVY

h2("4.2 What each layer does")
make_table(
    ["Layer", "Responsibility"],
    [
        ["1. Presentation", "What the engineer sees and uses: the chat window, the equipment-health dashboard (charts, trends, anomalies), alert notifications, downloadable reports, and role-based views."],
        ["2. Agentic orchestration", "The decision-making brain. The orchestrator plans and calls the specialised agents, who use tools, then composes a single explainable answer."],
        ["3. Knowledge (RAG)", "Stores all documents as searchable \"embeddings\" in a vector database so the right manual/SOP/past-case can be retrieved instantly and cited."],
        ["4. Predictive / ML engine", "Math + ML models for anomaly detection, Remaining Useful Life, and failure prediction on sensor data."],
        ["5. Data layer", "Where everything lives: incoming (simulated) sensor stream, logs, historical records, spares catalogue, feedback, and the digital logbook."],
    ],
    widths=[1.7, 4.9],
)

callout("Why this design wins points:",
        "It directly maps every functional requirement to a clear component, it is explainable (RAG citations + "
        "visible agent reasoning), it is scalable (agents and data sources can be added independently), and it "
        "demonstrates genuine agentic behaviour — exactly the judging criteria.")

# =========================================================
# 5. TECH STACK
# =========================================================
h1("5. Technology Stack")
body("We choose tools that are free / open-source where possible, fast to build with in a hackathon, and "
     "production-credible. Public LLM APIs are explicitly allowed by the brief.")

make_table(
    ["Area", "Choice", "Why"],
    [
        ["Language", "Python 3", "Best ecosystem for AI/ML, RAG, and data work."],
        ["LLM (reasoning)", "OpenAI GPT (API) — chat + embeddings", "Strong reasoning for every agent; one LLM, no SLM. API use is allowed by the brief."],
        ["Agent framework", "Lightweight custom orchestrator (Python)", "Transparent multi-agent planning, tool-calling, and memory; no heavy dependency."],
        ["RAG / Vector store", "Self-contained NumPy cosine store + OpenAI embeddings", "Fast, local, zero dependency risk on Python 3.8; drop-in replaceable with Chroma/FAISS."],
        ["Predictive ML", "scikit-learn, statsmodels, (optional) PyTorch", "Anomaly detection (Isolation Forest), RUL & failure models."],
        ["Backend / API", "FastAPI", "Lightweight, async, auto-docs; serves REST + WebSocket alerts."],
        ["Frontend", "Streamlit — PRIMARY", "Gets the chat + health dashboard demo up fast within hackathon time (React only if time allows)."],
        ["Dashboard / charts", "Plotly", "Interactive health trends, anomaly and risk visualisations."],
        ["Data stores", "SQLite/PostgreSQL + JSON/CSV", "History, spares, feedback, and logbook persistence."],
        ["IoT simulation", "Python data generator / MQTT (optional)", "Streams synthetic sensor data to mimic live equipment."],
        ["Packaging", "Docker + requirements.txt", "One-command, reproducible run for the judges."],
    ],
    widths=[1.5, 2.4, 2.7],
)

callout("Committed (as-built) stack:",
        "Streamlit + OpenAI GPT + a self-contained NumPy vector store + scikit-learn. This combination "
        "demonstrates ALL features (including every optional enhancement) and runs with zero dependency "
        "friction on the target Python 3.8 environment.")

# =========================================================
# 6. DATA & SYSTEM FLOW
# =========================================================
h1("6. Data Flow & System Flow")
numbered("Documents (manuals, SOPs, history) are ingested once, split into chunks, embedded, and stored in the vector DB.", bold_lead="Ingestion — ")
numbered("Simulated sensor data streams continuously into the data layer; the Anomaly agent watches it.", bold_lead="Live data — ")
numbered("The engineer asks a question (or an alert fires automatically).", bold_lead="Trigger — ")
numbered("The Orchestrator plans and dispatches the relevant specialist agents.", bold_lead="Planning — ")
numbered("Agents retrieve documents (RAG) and run ML models as needed.", bold_lead="Tool use — ")
numbered("Results are merged into a structured, explainable answer with citations.", bold_lead="Synthesis — ")
numbered("Reports + logbook entries are generated; role-based alerts are sent.", bold_lead="Output — ")
numbered("Engineer feedback is stored and used to improve future answers.", bold_lead="Learning — ")

# =========================================================
# 7. PLAN / ROADMAP
# =========================================================
h1("7. Build Plan (Phased)")
make_table(
    ["Phase", "What we deliver"],
    [
        ["Phase 1 — Foundation", "Sample dataset (manuals, logs, sensor CSVs, spares), RAG pipeline, basic chat that answers from documents with citations."],
        ["Phase 2 — Intelligence", "Diagnostic, Root-Cause & Recommendation agents; anomaly detection + RUL/failure prediction models."],
        ["Phase 3 — Orchestration", "Full multi-agent orchestrator, risk & priority scoring, multi-turn memory."],
        ["Phase 4 — Experience", "Dashboard, real-time alerts, reports, digital logbook, role-based views, feedback loop."],
        ["Phase 5 — Polish & deliver", "Dockerise, write the architecture doc, prepare sample I/O, record the demo video."],
    ],
    widths=[1.8, 4.8],
)

# =========================================================
# 8. DELIVERABLES
# =========================================================
h1("8. Deliverables Checklist (from the brief)")
bullet("Detailed source code of the working prototype (everything needed to run it).")
bullet("A document covering: architecture, tech stack, data & system flow, model/reasoning pipeline, alerting & prediction logic, assumptions & limitations.")
bullet("Install / configure / run instructions.")
bullet("Sample input & output demonstration.")
bullet("A screen-recording video showcasing the built features.")

# =========================================================
# 9. ASSUMPTIONS & LIMITATIONS
# =========================================================
h1("9. Assumptions & Limitations")
bullet("We use realistic synthetic / sample data (manuals, logs, sensor streams), since live plant data is not available.")
bullet("IoT input is simulated; the architecture is ready to connect to real sensors/MQTT later.")
bullet("LLM quality depends on the API/model used; the RAG grounding and citations reduce hallucination but do not eliminate it.")
bullet("Predictive models are trained on limited sample data and demonstrate the method, not plant-certified accuracy.")
bullet("The prototype is a decision-support tool — a human engineer always makes the final call.")

# =========================================================
# 10. GLOSSARY
# =========================================================
h1("10. Glossary (Plain-English)")
make_table(
    ["Term", "Meaning"],
    [
        ["Agentic AI", "AI that can plan and carry out multi-step tasks using tools and decisions, not just answer once."],
        ["LLM / SLM", "Large / Small Language Model — the AI that understands and generates text."],
        ["RAG", "Retrieval-Augmented Generation — fetch relevant documents first, then answer using them (so answers are grounded and citable)."],
        ["Embedding / Vector DB", "Turning text into numbers so the computer can find the most relevant passages quickly."],
        ["Anomaly detection", "Spotting sensor readings that are abnormal compared to normal behaviour."],
        ["RUL", "Remaining Useful Life — an estimate of how long equipment will keep working before failure."],
        ["Root-cause analysis", "Finding the underlying reason a failure happened, not just the symptom."],
        ["Orchestrator", "The manager agent that coordinates all the specialist agents."],
        ["SOP", "Standard Operating Procedure — official approved steps for a task."],
        ["Lead time", "How long it takes for an ordered spare part to actually arrive."],
    ],
    widths=[1.7, 4.9],
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— End of Understanding & Plan Document —")
r.italic = True
r.font.color.rgb = GREY
r.font.size = Pt(10)

# ---------- Footer ----------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Tata Steel AI Hackathon 2026  •  Maintenance Wizard  •  Round 2: Agentic AI Challenge")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY

out = "/home/hackerearth331/Desktop/Tata Steel Hackathon/Maintenance_Wizard_Understanding_and_Plan.docx"
doc.save(out)
print("SAVED:", out)
